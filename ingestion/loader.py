"""Document loading utilities for local file ingestion.

This module provides :class:`Loader` which loads text from markdown, text, and
PDF files as well as MITRE ATT&CK JSON data for indexing.
"""

import os
import re
import zipfile
from typing import List, Tuple, Dict, Any
from pypdf import PdfReader
from docx import Document
from mitre_chunker import load_mitre_documents

class Loader:
    """Load documents from disk for ingestion into retrieval indexes."""

    def load_documents(self, folder: str) -> List[str]:
        """Load all supported documents from a folder.

        Args:
            folder (str): Path to the directory containing source files.

        Returns:
            List[str]: A list of loaded document texts.
        """
        documents = []
        try:
            files = os.listdir(folder)
        except OSError as e:
            print(f'Error reading folder {folder}: {e}')
            return []
        for filename in files:
            file_path = os.path.join(folder, filename)
            try:
                if filename.endswith(('.txt', '.md')):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        documents.append(content)
                elif filename.endswith('.pdf'):
                    content = self._extract_pdf_text(file_path)
                    if content:
                        documents.append(content)
                elif filename.endswith('.docx'):
                    content = self._extract_docx_text(file_path)
                    if content:
                        documents.append(content)
            except Exception as e:
                print(f'Error loading {filename}: {e}')
                continue
        return documents

    def _extract_pdf_text(self, pdf_path: str) -> str:
        try:
            reader = PdfReader(pdf_path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return '\n'.join(text_parts)
        except Exception as e:
            print(f'Error extracting PDF text from {pdf_path}: {e}')
            return ''

    def _extract_docx_text(self, docx_path: str) -> str:
        try:
            document = Document(docx_path)
            text_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            for table in document.tables:
                for row in table.rows:
                    row_text = ' '.join(cell.text for cell in row.cells if cell.text)
                    if row_text:
                        text_parts.append(row_text)
            return '\n'.join(text_parts)
        except Exception as e:
            print(f'Error extracting DOCX text from {docx_path}: {e}')
            return self._extract_docx_text_fallback(docx_path)

    def _extract_docx_text_fallback(self, docx_path: str) -> str:
        try:
            with zipfile.ZipFile(docx_path, 'r') as z:
                with z.open('word/document.xml') as doc_xml:
                    content = doc_xml.read().decode('utf-8')
        except Exception as e:
            print(f'Fallback extraction failed for {docx_path}: {e}')
            return ''

        text_parts = re.findall(r'<w:t[^>]*>(.*?)</w:t>', content)
        if not text_parts:
            return ''

        return '\n'.join(text_parts)

    def load_mitre(self, json_path: str) -> Tuple[List[str], List[Dict[str, Any]]]:
        try:
            documents = load_mitre_documents(json_path)
            texts = [doc["text"] for doc in documents]
            metadata = [doc["metadata"] for doc in documents]
            return (texts, metadata)
        except Exception as e:
            print(f'Error loading MITRE data from {json_path}: {e}')
            return ([], [])